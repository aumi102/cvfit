from __future__ import annotations

import argparse
import html
import json
import os
import re
import sys
import tempfile
import time
import uuid
import zipfile
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import urlencode, urljoin
from urllib.request import Request, urlopen
from xml.etree import ElementTree


DEFAULT_TIMEOUT_SECONDS = 300
REQUEST_TIMEOUT_SECONDS = 30
POLL_INTERVAL_SECONDS = 3
TERMINAL_STATUSES = {"succeeded", "failed"}
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOCX_V3_SECTIONS = (
    "Improvement Action Plan",
    "Safe Rewrite Suggestions",
    "Interview Prep",
    "Learning Roadmap",
    "Limitations / Safety Notes",
)
V3_RESULT_FIELDS = (
    "improvement_actions",
    "safe_rewrite_suggestions",
    "interview_prep",
    "learning_roadmap",
    "limitations",
)
COMPARISON_REQUIRED_KEYS = (
    "base_job_id",
    "new_job_id",
    "base_score",
    "new_score",
    "score_delta",
    "breakdown_delta",
    "resolved_missing_skills",
    "still_missing_skills",
    "newly_matched_skills",
    "new_evidence",
    "keyword_stuffing_warnings",
    "improvement_summary",
    "next_actions",
)
SENSITIVE_TERMS = (
    "access_token",
    "access_token_hash",
    "storage_path",
    "report_docx_path",
    "local_path",
    "s3_key",
    "object_key",
    "bucket",
    "password",
    "raw_cv_text",
    "cv_text",
)
REDACTION_PATTERNS = (
    re.compile(r"access_token=([^&\s]+)", re.IGNORECASE),
    re.compile(r"bearer\s+[A-Za-z0-9._\-]+", re.IGNORECASE),
    re.compile(r"jwt\s*[:=]\s*[A-Za-z0-9._\-]+", re.IGNORECASE),
    re.compile(r"https?://[^\s?]+?\?[^\s]+", re.IGNORECASE),
)
HIRING_GUARANTEE_PATTERNS = (
    re.compile(r"guarantee[sd]?\s+(an\s+)?interview", re.IGNORECASE),
    re.compile(r"guarantee[sd]?\s+(you\s+)?(will\s+)?(be\s+)?hired", re.IGNORECASE),
    re.compile(r"you\s+will\s+get\s+(the\s+)?job", re.IGNORECASE),
)


class SmokeError(RuntimeError):
    pass


def normalize_base_url(value: str) -> str:
    cleaned = value.strip()
    if not cleaned:
        raise SmokeError("API_BASE_URL is required, for example https://your-render-api.onrender.com")
    if not cleaned.startswith(("http://", "https://")):
        raise SmokeError("API_BASE_URL must start with http:// or https://")
    return cleaned.rstrip("/")


def env_flag_enabled(name: str) -> bool:
    return os.environ.get(name, "").strip() == "1"


def configured_timeout() -> int:
    raw_value = os.environ.get("SMOKE_TIMEOUT_SECONDS", str(DEFAULT_TIMEOUT_SECONDS)).strip()
    try:
        timeout = int(raw_value)
    except ValueError as exc:
        raise SmokeError("SMOKE_TIMEOUT_SECONDS must be an integer") from exc
    if timeout <= 0:
        raise SmokeError("SMOKE_TIMEOUT_SECONDS must be greater than zero")
    return timeout


def require_mutating_allowed() -> None:
    if not env_flag_enabled("SMOKE_ALLOW_MUTATING"):
        raise SmokeError("refusing mutating Phase 4 smoke; set SMOKE_ALLOW_MUTATING=1")


def redact_text(value: object) -> str:
    text = str(value)
    text = re.sub(
        r"(access_token[\"']?\s*[:=]\s*[\"']?)[^\"',\s&}]+",
        r"\1<hidden>",
        text,
        flags=re.IGNORECASE,
    )
    for pattern in REDACTION_PATTERNS:
        text = pattern.sub(_redacted_match, text)
    return text


def _redacted_match(match: re.Match[str]) -> str:
    text = match.group(0)
    if text.lower().startswith("access_token="):
        return "access_token=<hidden>"
    if text.lower().startswith("bearer "):
        return "Bearer <hidden>"
    if text.lower().startswith("jwt"):
        return "JWT=<hidden>"
    return "<redacted-url>"


def build_url(base_url: str, path_or_url: str) -> str:
    return urljoin(f"{base_url}/", path_or_url)


def request_json(
    base_url: str,
    method: str,
    path: str,
    body: dict | None = None,
    timeout: int = REQUEST_TIMEOUT_SECONDS,
) -> dict:
    data = None
    headers = {}
    if body is not None:
        data = json.dumps(body).encode("utf-8")
        headers["Content-Type"] = "application/json"

    req = Request(build_url(base_url, path), data=data, headers=headers, method=method)
    with urlopen(req, timeout=timeout) as response:
        return json.loads(response.read().decode("utf-8"))


def request_bytes(
    base_url: str,
    method: str,
    path_or_url: str,
    timeout: int = 60,
) -> tuple[bytes, dict[str, str]]:
    req = Request(build_url(base_url, path_or_url), method=method)
    with urlopen(req, timeout=timeout) as response:
        headers = {key.lower(): value for key, value in response.headers.items()}
        return response.read(), headers


def post_multipart(
    base_url: str,
    path: str,
    *,
    fields: dict[str, str] | None = None,
    file_field: str,
    file_path: Path,
    content_type: str,
) -> dict:
    boundary = f"----cvfit-phase4-smoke-{uuid.uuid4().hex}"
    parts: list[bytes] = []
    for name, value in (fields or {}).items():
        parts.extend(
            [
                f"--{boundary}\r\n".encode("utf-8"),
                f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode("utf-8"),
                value.encode("utf-8"),
                b"\r\n",
            ]
        )

    parts.extend(
        [
            f"--{boundary}\r\n".encode("utf-8"),
            (
                f'Content-Disposition: form-data; name="{file_field}"; '
                f'filename="{file_path.name}"\r\n'
            ).encode("utf-8"),
            f"Content-Type: {content_type}\r\n\r\n".encode("utf-8"),
            file_path.read_bytes(),
            b"\r\n",
            f"--{boundary}--\r\n".encode("utf-8"),
        ]
    )
    req = Request(
        build_url(base_url, path),
        data=b"".join(parts),
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    with urlopen(req, timeout=60) as response:
        return json.loads(response.read().decode("utf-8"))


def create_synthetic_docx_cv(*, revised: bool = False) -> Path:
    if revised:
        paragraphs = [
            "CVFit Phase 4 Revised Smoke Candidate",
            "Backend engineer using Python, FastAPI, PostgreSQL, Redis, Docker, Kubernetes, and SQL.",
            "Built FastAPI services with PostgreSQL and Redis-backed background workers.",
            "Containerized services with Docker and deployed a small Kubernetes workload with health checks.",
            "Designed SQL schemas, wrote tests, and documented deployment troubleshooting steps.",
            "Skills: Python, FastAPI, PostgreSQL, Redis, Docker, Kubernetes, SQL",
        ]
        prefix = "cvfit-phase4-revised-"
    else:
        paragraphs = [
            "CVFit Phase 4 Initial Smoke Candidate",
            "Backend engineer using Python, FastAPI, PostgreSQL, Docker, Redis, and SQL.",
            "Built FastAPI services with PostgreSQL and Redis-backed background workers.",
            "Implemented Docker deployment checks and API tests for internal services.",
            "Skills: Python, FastAPI, PostgreSQL, Redis, Docker, SQL",
        ]
        prefix = "cvfit-phase4-initial-"

    tmp = tempfile.NamedTemporaryFile(prefix=prefix, suffix=".docx", delete=False)
    tmp.close()
    path = Path(tmp.name)
    write_minimal_docx(path, paragraphs)
    return path


def write_minimal_docx(path: Path, paragraphs: list[str]) -> None:
    body = "".join(f"<w:p><w:r><w:t>{html.escape(item)}</w:t></w:r></w:p>" for item in paragraphs)
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{body}<w:sectPr/></w:body></w:document>"
    )
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as docx:
        docx.writestr(
            "[Content_Types].xml",
            (
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Override PartName="/word/document.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                "</Types>"
            ),
        )
        docx.writestr(
            "_rels/.rels",
            (
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" '
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
                'Target="word/document.xml"/>'
                "</Relationships>"
            ),
        )
        docx.writestr("word/document.xml", document_xml)


def extract_docx_text(docx_bytes: bytes) -> str | None:
    try:
        from docx import Document  # type: ignore
    except ModuleNotFoundError:
        return extract_docx_text_stdlib(docx_bytes)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_path = Path(tmp.name)
    try:
        doc = Document(str(tmp_path))
        parts = [paragraph.text for paragraph in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    finally:
        try:
            tmp_path.unlink()
        except FileNotFoundError:
            pass


def extract_docx_text_stdlib(docx_bytes: bytes) -> str | None:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_path = Path(tmp.name)
    try:
        with zipfile.ZipFile(tmp_path) as docx:
            xml_bytes = docx.read("word/document.xml")
    except (KeyError, zipfile.BadZipFile):
        return None
    finally:
        try:
            tmp_path.unlink()
        except FileNotFoundError:
            pass

    root = ElementTree.fromstring(xml_bytes)
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    texts = [node.text or "" for node in root.findall(".//w:t", namespace)]
    return "\n".join(texts)


def validate_result_v3_payload(payload: dict) -> tuple[str, float | int]:
    nested = payload.get("result") if isinstance(payload.get("result"), dict) else {}
    if not nested:
        raise SmokeError("result response missing result object")
    schema_version = nested.get("schema_version")
    if schema_version != "3.0":
        raise SmokeError(f"expected result.schema_version=3.0, got {schema_version!r}")

    scores = nested.get("scores") if isinstance(nested.get("scores"), dict) else {}
    overall = nested.get("overall") if isinstance(nested.get("overall"), dict) else {}
    fit_score = nested.get("fit_score")
    if fit_score is None:
        fit_score = payload.get("overall_fit_score")
    if fit_score is None:
        raise SmokeError("result missing overall_fit_score/result.fit_score")
    if scores.get("fit_score") is None:
        raise SmokeError("result missing result.scores.fit_score")
    if overall.get("fit_score") is None:
        raise SmokeError("result missing result.overall.fit_score")

    missing = [key for key in V3_RESULT_FIELDS if not isinstance(nested.get(key), list)]
    if missing:
        raise SmokeError(f"result v3 missing list fields: {', '.join(missing)}")
    return str(schema_version), fit_score


def validate_reanalysis_response(payload: dict, parent_job_id: str) -> tuple[str, str, int]:
    job_id = payload.get("job_id")
    access_token = payload.get("access_token")
    analysis_group_id = payload.get("analysis_group_id")
    revision_number = payload.get("revision_number")
    if not job_id or not access_token or not analysis_group_id:
        raise SmokeError("reanalyze response missing job_id, access_token, or analysis_group_id")
    if payload.get("parent_job_id") != parent_job_id:
        raise SmokeError("reanalyze response parent_job_id does not match initial job")
    try:
        revision = int(revision_number)
    except (TypeError, ValueError) as exc:
        raise SmokeError("reanalyze response revision_number is not an integer") from exc
    if revision < 2:
        raise SmokeError(f"reanalyze response revision_number must be >= 2, got {revision}")
    if payload.get("status") not in {None, "queued", "running", "accepted"}:
        raise SmokeError(f"unexpected reanalysis status: {payload.get('status')}")
    return str(job_id), str(access_token), revision


def validate_comparison_payload(payload: dict) -> float | int | None:
    missing = [key for key in COMPARISON_REQUIRED_KEYS if key not in payload]
    if missing:
        raise SmokeError(f"comparison response missing keys: {', '.join(missing)}")
    if not isinstance(payload.get("breakdown_delta"), dict):
        raise SmokeError("comparison breakdown_delta must be an object")
    for key in (
        "resolved_missing_skills",
        "still_missing_skills",
        "newly_matched_skills",
        "new_evidence",
        "keyword_stuffing_warnings",
        "next_actions",
    ):
        if not isinstance(payload.get(key), list):
            raise SmokeError(f"comparison {key} must be a list")
    validate_no_sensitive_payload(payload, label="comparison response")
    text = json.dumps(payload, sort_keys=True)
    for pattern in HIRING_GUARANTEE_PATTERNS:
        if pattern.search(text):
            raise SmokeError("comparison response contains hiring guarantee wording")
    return payload.get("score_delta")


def validate_no_sensitive_payload(payload: object, *, label: str) -> None:
    text = json.dumps(payload, sort_keys=True, default=str).lower()
    leaked = [term for term in SENSITIVE_TERMS if term.lower() in text]
    if leaked:
        raise SmokeError(f"{label} exposed sensitive terms: {', '.join(leaked)}")


def validate_docx_report(docx_bytes: bytes, headers: dict[str, str]) -> bool:
    content_type = headers.get("content-type", "")
    disposition = headers.get("content-disposition", "")
    if DOCX_CONTENT_TYPE not in content_type:
        raise SmokeError(f"unexpected DOCX content type: {content_type}")
    if "cvfit_report_" not in disposition and "attachment" not in disposition.lower():
        raise SmokeError(f"unexpected report content-disposition: {disposition}")
    if len(docx_bytes) < 1000:
        raise SmokeError(f"downloaded DOCX too small: {len(docx_bytes)} bytes")

    text = extract_docx_text(docx_bytes)
    if text is None:
        print("DOCX text parse skipped; byte and header checks passed")
        return False
    missing_sections = [section for section in DOCX_V3_SECTIONS if section not in text]
    if missing_sections:
        raise SmokeError(f"DOCX missing v3 sections: {', '.join(missing_sections)}")
    validate_no_sensitive_payload({"docx_text": text}, label="DOCX report")
    return True


def health_check(base_url: str) -> None:
    health = request_json(base_url, "GET", "/health")
    if health != {"status": "ok"}:
        raise SmokeError(f"unexpected health response: {health}")
    print("health ok")


def upload_cv(base_url: str, cv_path: Path) -> str:
    uploaded = post_multipart(
        base_url,
        "/v1/cv/upload",
        file_field="file",
        file_path=cv_path,
        content_type=DOCX_CONTENT_TYPE,
    )
    cv_file_id = uploaded.get("cv_file_id")
    if not cv_file_id:
        raise SmokeError(f"upload response missing cv_file_id: {redact_text(uploaded)}")
    return str(cv_file_id)


def create_score_job(base_url: str, cv_file_id: str) -> tuple[str, str]:
    jd_text = (
        "Backend Engineer role requiring Python, FastAPI, PostgreSQL, Docker, Kubernetes, "
        "SQL, API testing, background worker troubleshooting, and deployment checks."
    )
    job = request_json(
        base_url,
        "POST",
        "/v1/jobs/create-score",
        {
            "cv_file_id": cv_file_id,
            "jd_text": jd_text,
            "options": {
                "target_role": "Backend Engineer",
                "language": "en",
                "strictness": "balanced",
                "output_formats": ["json", "docx"],
            },
        },
    )
    job_id = job.get("job_id")
    access_token = job.get("access_token")
    if not job_id:
        raise SmokeError(f"create-score response missing job_id: {redact_text(job)}")
    if not access_token:
        raise SmokeError("create-score response missing access_token")
    return str(job_id), str(access_token)


def wait_for_job(base_url: str, job_id: str, timeout_seconds: int) -> dict:
    deadline = time.monotonic() + timeout_seconds
    status = {}
    while time.monotonic() < deadline:
        status = request_json(base_url, "GET", f"/v1/jobs/{job_id}")
        job_status = status.get("status")
        print(f"job_id={job_id} status={job_status} progress={status.get('progress')}")
        if job_status in TERMINAL_STATUSES:
            break
        time.sleep(POLL_INTERVAL_SECONDS)
    else:
        raise SmokeError(f"job {job_id} timed out after {timeout_seconds}s")

    if status.get("status") == "failed":
        raise SmokeError(f"job {job_id} failed: {status.get('error_message')}")
    if status.get("status") != "succeeded":
        raise SmokeError(f"unexpected terminal status for {job_id}: {status}")
    return status


def fetch_result(base_url: str, job_id: str, access_token: str) -> dict:
    token_query = urlencode({"access_token": access_token})
    return request_json(base_url, "GET", f"/v1/jobs/{job_id}/result?{token_query}")


def download_report(base_url: str, job_id: str, access_token: str) -> tuple[bytes, dict[str, str]]:
    token_query = urlencode({"access_token": access_token})
    report = request_json(base_url, "GET", f"/v1/jobs/{job_id}/report?{token_query}")
    if "local_path" in report:
        raise SmokeError("report metadata exposed local_path")
    if report.get("format") != "docx" or not isinstance(report.get("download_url"), str):
        raise SmokeError(f"unexpected report metadata: {redact_text(report)}")
    return request_bytes(base_url, "GET", report["download_url"])


def reanalyze_job(base_url: str, parent_job_id: str, parent_access_token: str, cv_path: Path) -> dict:
    return post_multipart(
        base_url,
        f"/v1/jobs/{parent_job_id}/reanalyze",
        fields={"access_token": parent_access_token},
        file_field="file",
        file_path=cv_path,
        content_type=DOCX_CONTENT_TYPE,
    )


def fetch_comparison(base_url: str, base_job_id: str, new_job_id: str, child_access_token: str) -> dict:
    token_query = urlencode({"access_token": child_access_token})
    return request_json(base_url, "GET", f"/v1/jobs/{base_job_id}/comparison/{new_job_id}?{token_query}")


def run_read_only(base_url: str) -> int:
    health_check(base_url)
    print("read-only Phase 4 smoke completed")
    print("mutating mode is required for full Phase 4 verification")
    print("mutating mode uploads two synthetic DOCX CVs, creates an initial job, reanalyzes it, downloads a DOCX report, and calls comparison")
    print("mutating mode leaves synthetic jobs/reports because this API has no cleanup endpoint")
    return 0


def run_mutating(base_url: str, timeout_seconds: int) -> int:
    initial_cv = create_synthetic_docx_cv(revised=False)
    revised_cv = create_synthetic_docx_cv(revised=True)
    summary: dict[str, object] = {
        "api_base_url": base_url,
        "docx_v3_sections_found": False,
    }
    try:
        health_check(base_url)

        cv_file_id = upload_cv(base_url, initial_cv)
        print(f"uploaded initial synthetic cv_file_id={cv_file_id}")

        initial_job_id, initial_token = create_score_job(base_url, cv_file_id)
        summary["initial_job_id"] = initial_job_id
        print(f"created initial job_id={initial_job_id}")

        wait_for_job(base_url, initial_job_id, timeout_seconds)
        initial_result = fetch_result(base_url, initial_job_id, initial_token)
        schema_version, fit_score = validate_result_v3_payload(initial_result)
        summary["schema_version"] = schema_version
        print(f"initial result v3 ok fit_score={fit_score}")

        docx_bytes, headers = download_report(base_url, initial_job_id, initial_token)
        docx_sections_found = validate_docx_report(docx_bytes, headers)
        summary["docx_v3_sections_found"] = docx_sections_found
        print(f"downloaded DOCX report bytes={len(docx_bytes)} v3_sections_found={docx_sections_found}")

        reanalysis = reanalyze_job(base_url, initial_job_id, initial_token, revised_cv)
        child_job_id, child_token, revision_number = validate_reanalysis_response(reanalysis, initial_job_id)
        summary["child_job_id"] = child_job_id
        summary["revision_number"] = revision_number
        print(f"created child job_id={child_job_id} revision_number={revision_number}")

        wait_for_job(base_url, child_job_id, timeout_seconds)
        child_result = fetch_result(base_url, child_job_id, child_token)
        child_schema_version, child_fit_score = validate_result_v3_payload(child_result)
        print(f"child result v3 ok schema_version={child_schema_version} fit_score={child_fit_score}")

        comparison = fetch_comparison(base_url, initial_job_id, child_job_id, child_token)
        score_delta = validate_comparison_payload(comparison)
        summary["comparison_score_delta"] = score_delta
        print(f"comparison ok score_delta={score_delta}")

        print("phase4 smoke summary:")
        for key, value in summary.items():
            print(f"- {key}: {value}")
        print("phase4 smoke test passed")
        return 0
    finally:
        for path in (initial_cv, revised_cv):
            try:
                path.unlink()
            except FileNotFoundError:
                pass


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Production Phase 4 smoke test")
    parser.add_argument(
        "--mutating",
        action="store_true",
        help="run the full synthetic Phase 4 flow; requires SMOKE_ALLOW_MUTATING=1",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(sys.argv[1:] if argv is None else argv)
    try:
        base_url = normalize_base_url(os.environ.get("API_BASE_URL", ""))
        if args.mutating:
            require_mutating_allowed()
            return run_mutating(base_url, configured_timeout())
        return run_read_only(base_url)
    except HTTPError as exc:
        body = redact_text(exc.read().decode("utf-8", errors="replace"))
        print(f"HTTP error {exc.code}: {body}", file=sys.stderr)
        return 1
    except (SmokeError, URLError, TimeoutError, OSError, ElementTree.ParseError) as exc:
        print(f"phase4 smoke failed: {redact_text(exc)}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
