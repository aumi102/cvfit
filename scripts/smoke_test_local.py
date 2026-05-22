from __future__ import annotations

import json
import os
import sys
import tempfile
import time
import uuid
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import urlencode
from urllib.request import Request, urlopen

from docx import Document


API_BASE_URL = os.environ.get("API_BASE_URL", "http://localhost:8000").rstrip("/")
TIMEOUT_SECONDS = int(os.environ.get("SMOKE_TIMEOUT_SECONDS", "300"))


def request_json(method: str, path: str, body: dict | None = None, headers: dict | None = None) -> dict:
    data = None
    final_headers = headers or {}
    if body is not None:
        data = json.dumps(body).encode("utf-8")
        final_headers = {"Content-Type": "application/json", **final_headers}

    req = Request(f"{API_BASE_URL}{path}", data=data, headers=final_headers, method=method)
    with urlopen(req, timeout=30) as response:
        return json.loads(response.read().decode("utf-8"))


def request_bytes(method: str, path: str) -> bytes:
    req = Request(f"{API_BASE_URL}{path}", method=method)
    with urlopen(req, timeout=60) as response:
        return response.read()


def expect_http_status(method: str, path: str, expected_status: int) -> None:
    req = Request(f"{API_BASE_URL}{path}", method=method)
    try:
        with urlopen(req, timeout=30) as response:
            status = response.status
    except HTTPError as exc:
        status = exc.code
        exc.read()
    if status != expected_status:
        raise RuntimeError(f"expected HTTP {expected_status} for {path}, got {status}")


def redact_report_metadata(report: dict) -> dict:
    redacted = dict(report)
    download_url = redacted.get("download_url")
    if isinstance(download_url, str) and "access_token=" in download_url:
        redacted["download_url"] = download_url.split("access_token=", 1)[0] + "access_token=<hidden>"
    return redacted


def post_multipart_file(path: str, field_name: str, file_path: Path, content_type: str) -> dict:
    boundary = f"----cvfit-smoke-{uuid.uuid4().hex}"
    file_bytes = file_path.read_bytes()
    parts = [
        f"--{boundary}\r\n".encode("utf-8"),
        (
            f'Content-Disposition: form-data; name="{field_name}"; '
            f'filename="{file_path.name}"\r\n'
        ).encode("utf-8"),
        f"Content-Type: {content_type}\r\n\r\n".encode("utf-8"),
        file_bytes,
        b"\r\n",
        f"--{boundary}--\r\n".encode("utf-8"),
    ]
    body = b"".join(parts)
    req = Request(
        f"{API_BASE_URL}{path}",
        data=body,
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    with urlopen(req, timeout=60) as response:
        return json.loads(response.read().decode("utf-8"))


def create_docx_cv() -> Path:
    doc = Document()
    doc.add_heading("CVFit Smoke Test Candidate", level=1)
    doc.add_paragraph("Email: smoke-test@example.com")
    doc.add_heading("Summary", level=2)
    doc.add_paragraph("Backend engineer with Python, FastAPI, PostgreSQL, Docker, and Redis experience.")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph(
        "Built FastAPI services with PostgreSQL and Redis queues for asynchronous background processing."
    )
    doc.add_paragraph(
        "Implemented Docker-based deployments and improved API latency by 25 percent."
    )
    doc.add_paragraph(
        "Designed SQL schemas, wrote API endpoints, and maintained production monitoring workflows."
    )
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Python, FastAPI, PostgreSQL, Redis, Docker, Git, SQL")

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    path = Path(tmp.name)
    doc.save(path)
    return path


def wait_for_health() -> None:
    deadline = time.time() + 60
    while time.time() < deadline:
        try:
            health = request_json("GET", "/health")
            if health.get("status") == "ok":
                print("health ok")
                return
        except (HTTPError, URLError, TimeoutError):
            time.sleep(2)
    raise RuntimeError("API health check did not become ready")


def main() -> int:
    cv_path = create_docx_cv()
    try:
        wait_for_health()

        uploaded = post_multipart_file(
            "/v1/cv/upload",
            "file",
            cv_path,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        cv_file_id = uploaded.get("cv_file_id")
        if not cv_file_id:
            raise RuntimeError(f"upload response missing cv_file_id: {uploaded}")
        print(f"uploaded cv_file_id={cv_file_id}")

        jd_text = """
        Backend Engineer role requiring Python, FastAPI, PostgreSQL, Redis, Docker, and SQL.
        The engineer will build API services, maintain background workers, improve performance,
        write tests, and support deployment workflows.
        """
        job = request_json(
            "POST",
            "/v1/jobs/create-score",
            {
                "cv_file_id": cv_file_id,
                "jd_text": jd_text,
                "options": {"output_formats": ["json", "docx"]},
            },
        )
        job_id = job.get("job_id")
        access_token = job.get("access_token")
        if not job_id:
            raise RuntimeError(f"create-score response missing job_id: {job}")
        if not access_token:
            raise RuntimeError(f"create-score response missing access_token: {job}")
        print(f"created job_id={job_id}")

        deadline = time.time() + TIMEOUT_SECONDS
        status = {}
        while time.time() < deadline:
            status = request_json("GET", f"/v1/jobs/{job_id}")
            print(f"job status={status.get('status')} progress={status.get('progress')}")
            if status.get("status") == "succeeded":
                break
            if status.get("status") == "failed":
                raise RuntimeError(f"job failed: {status.get('error_message')}")
            time.sleep(3)
        else:
            raise RuntimeError(f"job timed out after {TIMEOUT_SECONDS}s: {status}")

        expect_http_status("GET", f"/v1/jobs/{job_id}/result", 403)
        expect_http_status("GET", f"/v1/jobs/{job_id}/result?access_token=wrong-token", 403)
        print("access token protection ok")

        token_query = urlencode({"access_token": access_token})
        result = request_json("GET", f"/v1/jobs/{job_id}/result?{token_query}")
        scores = result.get("result", {}).get("scores", {})
        if "fit_score" not in scores:
            raise RuntimeError(f"result missing scores.fit_score: {result}")
        print(f"fit_score={scores['fit_score']}")

        report = request_json("GET", f"/v1/jobs/{job_id}/report?{token_query}")
        if "local_path" in report:
            raise RuntimeError(f"report metadata exposed local_path: {redact_report_metadata(report)}")
        if report.get("format") != "docx" or not report.get("download_url"):
            raise RuntimeError(f"unexpected report metadata: {redact_report_metadata(report)}")
        print(f"report metadata ok: {redact_report_metadata(report)}")

        report_bytes = request_bytes("GET", report["download_url"])
        if len(report_bytes) < 1000:
            raise RuntimeError(f"downloaded DOCX too small: {len(report_bytes)} bytes")
        print(f"downloaded report bytes={len(report_bytes)}")
        print("smoke test passed")
        return 0
    except HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        print(f"HTTP error {exc.code}: {body}", file=sys.stderr)
        return 1
    except Exception as exc:
        print(f"smoke test failed: {exc}", file=sys.stderr)
        return 1
    finally:
        try:
            cv_path.unlink()
        except FileNotFoundError:
            pass


if __name__ == "__main__":
    raise SystemExit(main())
