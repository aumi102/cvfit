from __future__ import annotations

import argparse
import json
import os
import sys
import tempfile
import time
import uuid
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import urlencode, urljoin
from urllib.request import Request, urlopen

from docx import Document


DEFAULT_TIMEOUT_SECONDS = 180
REQUEST_TIMEOUT_SECONDS = 30
POLL_INTERVAL_SECONDS = 3
TERMINAL_STATUSES = {"succeeded", "failed"}


class SmokeError(RuntimeError):
    pass


def normalize_base_url(value: str) -> str:
    cleaned = value.strip()
    if not cleaned:
        raise SmokeError("API_BASE_URL is required, for example https://your-render-api.onrender.com")
    if not cleaned.startswith(("http://", "https://")):
        raise SmokeError("API_BASE_URL must start with http:// or https://")
    return cleaned.rstrip("/")


def env_flag_enabled(name: str) -> bool:
    return os.environ.get(name, "").strip() == "1"


def configured_timeout() -> int:
    raw_value = os.environ.get("SMOKE_TIMEOUT_SECONDS", str(DEFAULT_TIMEOUT_SECONDS)).strip()
    try:
        timeout = int(raw_value)
    except ValueError as exc:
        raise SmokeError("SMOKE_TIMEOUT_SECONDS must be an integer") from exc
    if timeout <= 0:
        raise SmokeError("SMOKE_TIMEOUT_SECONDS must be greater than zero")
    return timeout


def build_url(base_url: str, path: str) -> str:
    return urljoin(f"{base_url}/", path.lstrip("/"))


def request_json(
    base_url: str,
    method: str,
    path: str,
    body: dict | None = None,
    timeout: int = REQUEST_TIMEOUT_SECONDS,
) -> dict:
    data = None
    headers = {}
    if body is not None:
        data = json.dumps(body).encode("utf-8")
        headers["Content-Type"] = "application/json"

    req = Request(build_url(base_url, path), data=data, headers=headers, method=method)
    with urlopen(req, timeout=timeout) as response:
        return json.loads(response.read().decode("utf-8"))


def request_bytes(
    base_url: str,
    method: str,
    path: str,
    timeout: int = 60,
) -> bytes:
    req = Request(build_url(base_url, path), method=method)
    with urlopen(req, timeout=timeout) as response:
        return response.read()


def post_multipart_file(
    base_url: str,
    path: str,
    field_name: str,
    file_path: Path,
    content_type: str,
) -> dict:
    boundary = f"----cvfit-smoke-{uuid.uuid4().hex}"
    file_bytes = file_path.read_bytes()
    parts = [
        f"--{boundary}\r\n".encode("utf-8"),
        (
            f'Content-Disposition: form-data; name="{field_name}"; '
            f'filename="{file_path.name}"\r\n'
        ).encode("utf-8"),
        f"Content-Type: {content_type}\r\n\r\n".encode("utf-8"),
        file_bytes,
        b"\r\n",
        f"--{boundary}--\r\n".encode("utf-8"),
    ]
    req = Request(
        build_url(base_url, path),
        data=b"".join(parts),
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    with urlopen(req, timeout=60) as response:
        return json.loads(response.read().decode("utf-8"))


def create_synthetic_docx_cv() -> Path:
    doc = Document()
    doc.add_heading("CVFit Synthetic Smoke Candidate", level=1)
    doc.add_paragraph("Synthetic profile for automated MVP smoke testing.")
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(
        "Backend engineer using Python, FastAPI, PostgreSQL, Redis, Docker, and SQL."
    )
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Built small API services, background workers, tests, and deployment checks.")
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Python, FastAPI, PostgreSQL, Redis, Docker, SQL")

    tmp = tempfile.NamedTemporaryFile(prefix="cvfit-smoke-", suffix=".docx", delete=False)
    tmp.close()
    path = Path(tmp.name)
    doc.save(path)
    return path


def redact_url(value: str) -> str:
    if "access_token=" not in value:
        return value
    return value.split("access_token=", 1)[0] + "access_token=<hidden>"


def redact_report_metadata(report: dict) -> dict:
    redacted = dict(report)
    download_url = redacted.get("download_url")
    if isinstance(download_url, str):
        redacted["download_url"] = redact_url(download_url)
    return redacted


def print_read_only_plan() -> None:
    print("read-only smoke completed")
    print("mutating smoke is skipped by default")
    print("mutating smoke would upload one tiny synthetic DOCX, create one score job, poll it, fetch result JSON, and download the DOCX report")
    print("mutating smoke leaves one synthetic job/report because this API has no cleanup endpoint")


def run_read_only(base_url: str) -> int:
    health = request_json(base_url, "GET", "/health")
    if health != {"status": "ok"}:
        raise SmokeError(f"unexpected health response: {health}")
    print("health ok")
    print_read_only_plan()
    return 0


def run_mutating(base_url: str, timeout_seconds: int) -> int:
    cv_path = create_synthetic_docx_cv()
    try:
        health = request_json(base_url, "GET", "/health")
        if health != {"status": "ok"}:
            raise SmokeError(f"unexpected health response: {health}")
        print("health ok")

        uploaded = post_multipart_file(
            base_url,
            "/v1/cv/upload",
            "file",
            cv_path,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        cv_file_id = uploaded.get("cv_file_id")
        if not cv_file_id:
            raise SmokeError(f"upload response missing cv_file_id: {uploaded}")
        print(f"uploaded synthetic cv_file_id={cv_file_id}")

        jd_text = (
            "Backend Engineer role requiring Python, FastAPI, PostgreSQL, Redis, Docker, "
            "SQL, API testing, deployment checks, and background worker troubleshooting."
        )
        job = request_json(
            base_url,
            "POST",
            "/v1/jobs/create-score",
            {
                "cv_file_id": cv_file_id,
                "jd_text": jd_text,
                "options": {
                    "target_role": "Backend Engineer",
                    "language": "en",
                    "strictness": "balanced",
                    "output_formats": ["json", "docx"],
                },
            },
        )
        job_id = job.get("job_id")
        access_token = job.get("access_token")
        if not job_id:
            raise SmokeError(f"create-score response missing job_id: {job}")
        if not access_token:
            raise SmokeError("create-score response missing access_token")
        print(f"created synthetic job_id={job_id}")

        deadline = time.monotonic() + timeout_seconds
        status = {}
        while time.monotonic() < deadline:
            status = request_json(base_url, "GET", f"/v1/jobs/{job_id}")
            job_status = status.get("status")
            print(f"job status={job_status} progress={status.get('progress')}")
            if job_status in TERMINAL_STATUSES:
                break
            time.sleep(POLL_INTERVAL_SECONDS)
        else:
            raise SmokeError(f"job timed out after {timeout_seconds}s: {status}")

        if status.get("status") == "failed":
            raise SmokeError(f"job failed: {status.get('error_message')}")
        if status.get("status") != "succeeded":
            raise SmokeError(f"unexpected terminal job status: {status}")

        token_query = urlencode({"access_token": access_token})
        result = request_json(base_url, "GET", f"/v1/jobs/{job_id}/result?{token_query}")
        scores = result.get("result", {}).get("scores", {})
        if "fit_score" not in scores:
            raise SmokeError(f"result missing scores.fit_score: {result}")
        print(f"fit_score={scores['fit_score']}")

        report = request_json(base_url, "GET", f"/v1/jobs/{job_id}/report?{token_query}")
        if "local_path" in report:
            raise SmokeError("report metadata exposed local_path")
        download_url = report.get("download_url")
        if report.get("format") != "docx" or not isinstance(download_url, str):
            raise SmokeError(f"unexpected report metadata: {redact_report_metadata(report)}")
        print(f"report metadata ok: {redact_url(download_url)}")

        report_bytes = request_bytes(base_url, "GET", download_url)
        if len(report_bytes) < 1000:
            raise SmokeError(f"downloaded DOCX too small: {len(report_bytes)} bytes")
        print(f"downloaded report bytes={len(report_bytes)}")
        print("mutating smoke created one synthetic job/report; no cleanup endpoint is available")
        print("mvp smoke test passed")
        return 0
    finally:
        try:
            cv_path.unlink()
        except FileNotFoundError:
            pass


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Production-safe MVP smoke test")
    parser.add_argument(
        "--mutating",
        action="store_true",
        help="run the synthetic upload/job/result/report flow; requires SMOKE_ALLOW_MUTATING=1",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(sys.argv[1:] if argv is None else argv)
    try:
        base_url = normalize_base_url(os.environ.get("API_BASE_URL", ""))
        if args.mutating:
            if not env_flag_enabled("SMOKE_ALLOW_MUTATING"):
                raise SmokeError("refusing mutating smoke; set SMOKE_ALLOW_MUTATING=1 and use synthetic data only")
            return run_mutating(base_url, configured_timeout())
        return run_read_only(base_url)
    except HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        print(f"HTTP error {exc.code}: {body}", file=sys.stderr)
        return 1
    except (SmokeError, URLError, TimeoutError) as exc:
        print(f"mvp smoke failed: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
