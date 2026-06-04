from __future__ import annotations

from docx import Document

from app.services.reporting.report_docx import build_docx_report


def read_docx_text(path):
    doc = Document(path)
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def v2_result():
    return {
        "schema_version": "2.0",
        "fit_score": 78.4,
        "scores": {
            "fit_score": 78.4,
            "skill_match": 80,
            "responsibility_match": 75,
            "experience_level": 70,
            "project_relevance": 82,
            "cv_quality": 90,
        },
        "overall": {
            "fit_score": 78.4,
            "fit_level": "good",
            "summary": "Good fit with a PostgreSQL evidence gap.",
            "guardrail_notice": "This analysis estimates CV-to-JD fit only and does not guarantee any hiring outcome.",
        },
        "score_breakdown": [
            {
                "key": "skill_match",
                "label": "Skill match",
                "score": 80,
                "weight": 0.35,
                "explanation": "Most required backend skills were found.",
            }
        ],
        "matched_skills": [
            {
                "skill": "FastAPI",
                "requirement_type": "must_have",
                "confidence": 0.9,
                "jd_requirement": "Required: FastAPI service development.",
                "cv_evidence_ids": ["ev_cv_skill_001"],
                "jd_evidence_ids": ["ev_jd_skill_must_001"],
            }
        ],
        "missing_skills": [
            {
                "skill": "PostgreSQL",
                "requirement_type": "must_have",
                "severity": "high",
                "jd_requirement": "Required: PostgreSQL database experience.",
                "reason": "JD requires PostgreSQL, but PostgreSQL evidence was not found in the parsed CV.",
                "jd_evidence_ids": ["ev_jd_skill_must_002"],
                "suggestion": (
                    "If you have actually used PostgreSQL, add a truthful CV bullet with project context. "
                    "Only add this if it is true."
                ),
            }
        ],
        "evidence": [
            {
                "id": "ev_cv_skill_001",
                "source": "cv",
                "kind": "skill",
                "text": "Built FastAPI services for internal users.",
                "normalized_skill": "FastAPI",
            },
            {
                "id": "ev_cv_resp_001",
                "source": "cv",
                "kind": "responsibility",
                "best_cv_bullet": "Built API services and background workers.",
                "jd_requirement": "Build APIs",
                "similarity": 0.8,
                "jd_evidence_ids": ["ev_jd_resp_001"],
            },
            {
                "id": "ev_jd_skill_must_001",
                "source": "jd",
                "kind": "requirement",
                "text": "Required: FastAPI service development.",
            },
            {
                "id": "ev_jd_resp_001",
                "source": "jd",
                "kind": "responsibility",
                "text": "Build APIs",
            },
        ],
        "improvement_actions": [
            {
                "id": "act_missing_must_001",
                "type": "skill_gap",
                "priority": "high",
                "title": "Address PostgreSQL evidence gap",
                "suggestion": (
                    "If you have actually used PostgreSQL, add a project bullet with measurable outcome. "
                    "Only add this if it is true."
                ),
                "related_skill": "PostgreSQL",
                "related_evidence_ids": ["ev_jd_skill_must_002"],
                "guardrail": "Only add this if it is true. Do not invent skills or experience.",
            }
        ],
        "limitations": [
            "This analysis estimates CV-to-JD fit only and does not guarantee any hiring outcome.",
            "Missing evidence is not absolute proof that the candidate lacks a skill.",
        ],
        "metadata": {
            "generated_at": "2026-06-04T09:30:00+00:00",
            "scorer_version": "phase3.result_json_v2",
            "cv": {
                "parsed_confidence": 0.85,
                "skills_detected": ["FastAPI", "Python"],
            },
            "jd": {"role": "Backend Engineer"},
        },
    }


def legacy_result():
    return {
        "scores": {
            "fit_score": 66,
            "skill_match": 60,
            "responsibility_match": 70,
        },
        "skills": {
            "matched_must_groups": [{"group": ["python"], "matched_by": "python"}],
            "matched_nice_groups": [],
        },
        "skill_gap": {
            "missing_must_have": ["FastAPI"],
            "missing_nice_to_have": ["Docker"],
            "learn_suggestions": [
                {
                    "skill": "FastAPI",
                    "reason": "Required but not sufficiently evidenced",
                    "resources_level": "beginner",
                }
            ],
        },
        "cv_improvements": [
            {
                "issue": "Low number of measurable achievements",
                "fix": "add truthful metrics",
            }
        ],
        "evidence": [
            {
                "type": "skill_match",
                "text": "Python backend project.",
            }
        ],
        "jd": {"role": "Backend Engineer"},
    }


def test_build_docx_report_works_with_v2_result(tmp_path):
    out_path = tmp_path / "report-v2.docx"

    build_docx_report(v2_result(), str(out_path))

    text = read_docx_text(out_path)
    assert "AI CV Fit Report" in text
    assert "Executive Summary" in text
    assert "Score Breakdown" in text
    assert "Matched Skills" in text
    assert "Missing Skills / Gaps" in text
    assert "Evidence Highlights" in text
    assert "Improvement Actions" in text
    assert "Limitations" in text


def test_build_docx_report_works_with_legacy_result(tmp_path):
    out_path = tmp_path / "report-legacy.docx"

    build_docx_report(legacy_result(), str(out_path))

    text = read_docx_text(out_path)
    assert "AI CV Fit Report" in text
    assert "Fit score: 66" in text
    assert "Fit level: partial" in text
    assert "python" in text
    assert "FastAPI" in text


def test_docx_report_contains_v2_content(tmp_path):
    out_path = tmp_path / "report-content.docx"

    build_docx_report(v2_result(), str(out_path))

    text = read_docx_text(out_path)
    assert "Fit level: good" in text
    assert "FastAPI" in text
    assert "JD requires PostgreSQL, but PostgreSQL evidence was not found in the parsed CV." in text
    assert "If you have actually used PostgreSQL" in text
    assert "Only add this if it is true" in text


def test_docx_report_excludes_sensitive_keys_and_values(tmp_path):
    result = v2_result()
    result.update(
        {
            "access_token": "secret-token",
            "access_token_hash": "secret-hash",
            "storage_path": "uploads/private-cv.docx",
            "report_docx_path": "reports/private.docx",
            "s3_key": "private/key.docx",
            "local_path": "C:/private/cv.docx",
            "raw_cv_text": "full private cv",
        }
    )
    result["evidence"].append(
        {
            "id": "ev_cv_skill_unsafe",
            "source": "cv",
            "kind": "skill",
            "text": "access_token=secret-token storage_path uploads/private-cv.docx local_path C:/private/cv.docx",
        }
    )
    out_path = tmp_path / "report-safe.docx"

    build_docx_report(result, str(out_path))

    text = read_docx_text(out_path)
    for unsafe in (
        "access_token",
        "access_token_hash",
        "storage_path",
        "report_docx_path",
        "s3_key",
        "local_path",
        "raw_cv_text",
        "secret-token",
        "secret-hash",
        "uploads/private-cv.docx",
        "reports/private.docx",
        "private/key.docx",
        "C:/private/cv.docx",
        "full private cv",
    ):
        assert unsafe not in text
    assert "guaranteed hired" not in text.lower()
    assert "guarantee hired" not in text.lower()
